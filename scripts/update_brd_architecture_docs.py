from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
BRD_PATH = ROOT / "camwatch_brd.docx"
ARCH_PATH = ROOT / "camwatch_architecture.docx"
BRD_OUT_PATH = ROOT / "camwatch_brd_updated.docx"
ARCH_OUT_PATH = ROOT / "camwatch_architecture_updated.docx"


def insert_after(paragraph, text: str, style: str | None = None):
    new_p_xml = OxmlElement("w:p")
    paragraph._p.addnext(new_p_xml)
    new_p = Paragraph(new_p_xml, paragraph._parent)
    new_p.text = text
    if style:
        new_p.style = style
    return new_p


def ensure_brd():
    doc = Document(BRD_PATH)
    paragraphs = doc.paragraphs
    by_text = {p.text.strip(): p for p in paragraphs if p.text.strip()}

    in_scope_anchor = by_text["Daily, weekly, and monthly uptime reporting"]
    p = insert_after(in_scope_anchor, "Google Single Sign-On (SSO) using approved company accounts")
    insert_after(p, "Role-based access control aligned to Admin and User application permissions")
    p = insert_after(p, "Email alert delivery through SendGrid with fallback logging when credentials are unavailable")

    out_scope_anchor = by_text["Enterprise messaging provider rollout as part of the base release"]
    insert_after(out_scope_anchor, "Advanced identity federation beyond Google SSO in the current phase")

    functional_anchor = by_text["6. Functional Requirements"]
    p = insert_after(functional_anchor, "6.1 Authentication and Access Control", style="Heading 2")
    p = insert_after(p, "The system shall support email/password login for local administrator access.", style=None)
    p = insert_after(p, "The system shall support Google SSO for approved users using company email identities.", style=None)
    p = insert_after(p, "The system shall map authenticated email addresses to existing CamWatch user records and apply stored roles without changing the RBAC model.", style=None)
    p = insert_after(p, "The system shall restrict create, edit, delete, import, and alert-management actions to Admin users only.", style=None)
    p = insert_after(p, "The system shall allow User-role accounts to view dashboard, sites, devices, alerts, and reports in read-only mode.", style=None)
    p = insert_after(p, "6.2 Notification and Email Delivery", style="Heading 2")
    p = insert_after(p, "The system shall send alert creation, escalation, recovery, and resolution emails through SendGrid when a valid API key and sender identity are configured.", style=None)
    p = insert_after(p, "The system shall continue to preserve provider abstraction so email delivery can fall back to logging without changing alert business logic.", style=None)
    p = insert_after(p, "The system shall maintain notification history for operational audit and troubleshooting.", style=None)

    assumptions_anchor = by_text["Dockerized deployment is the primary hosting model for the current solution."]
    p = insert_after(assumptions_anchor, "Google Cloud OAuth configuration and approved client IDs are available for SSO enablement.")
    insert_after(p, "A verified sender identity and API credentials are available in SendGrid for production email delivery.")

    success_anchor = by_text["Reports support operational review of uptime and recurring failures."]
    p = insert_after(success_anchor, "Authorized users can sign in through Google SSO and immediately receive their correct application role.")
    insert_after(p, "Alert emails are delivered through SendGrid and recorded in notification history for traceability.")

    doc.save(BRD_OUT_PATH)


def ensure_architecture():
    doc = Document(ARCH_PATH)
    paragraphs = doc.paragraphs
    by_text = {p.text.strip(): p for p in paragraphs if p.text.strip()}

    backend_anchor = by_text["Background scheduling runs in-process with APScheduler at startup."]
    p = insert_after(backend_anchor, "3.1 Authentication Architecture", style="Heading 2")
    p = insert_after(p, "Local authentication uses FastAPI password login and JWT access tokens for session continuity.", style=None)
    p = insert_after(p, "Google SSO is integrated as an identity-provider entry path: the frontend obtains a Google credential, the backend verifies it against Google, then CamWatch issues its own JWT.", style=None)
    p = insert_after(p, "Role assignment remains internal to CamWatch by resolving the authenticated email address to an existing user record and applying the stored ADMIN or USER role.", style=None)

    notification_anchor = by_text["8. Notification Framework"]
    p = insert_after(notification_anchor, "Primary production email delivery is implemented through SendGrid over its HTTP API.", style=None)
    p = insert_after(p, "If SendGrid credentials are unavailable, the notification layer falls back to logged delivery intent while still writing notification history records.", style=None)
    p = insert_after(p, "Channel abstraction remains intact so WhatsApp, SMS, Teams, and Slack can continue to use the same send_notification contract.", style=None)
    p = insert_after(p, "8.1 Email Delivery Flow", style="Heading 2")
    p = insert_after(p, "Alert services prepare subject, message body, and recipient list from global configuration plus site-level contacts.", style=None)
    p = insert_after(p, "The notification provider chooses SendGrid when configured, otherwise SMTP or logging fallback depending on environment configuration.", style=None)
    p = insert_after(p, "Each delivery attempt is written to notification history for audit, operator visibility, and troubleshooting.", style=None)

    deployment_anchor = by_text["10. Deployment Architecture"]
    p = insert_after(deployment_anchor, "Runtime configuration now includes monitoring settings, JWT secrets, SendGrid credentials, and Google SSO client metadata supplied through environment variables.", style=None)
    insert_after(p, "This keeps deployment portable across Docker-based local, staging, and production environments without changing application code.", style=None)

    ops_anchor = by_text["Data hygiene in imports matters because duplicate or stale records distort site counts and reports."]
    p = insert_after(ops_anchor, "Google SSO depends on the configured client ID matching the frontend origin and on approved users already existing in the CamWatch user table.", style=None)
    insert_after(p, "SendGrid delivery depends on verified sender setup, valid API keys, and outbound internet access from the backend container.")

    doc.save(ARCH_OUT_PATH)


if __name__ == "__main__":
    ensure_brd()
    ensure_architecture()
